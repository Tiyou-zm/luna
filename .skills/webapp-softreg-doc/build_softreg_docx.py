from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def clean_inline(text: str) -> str:
    return text.replace("`", "").strip()


def set_run_font(run, ascii_font: str, east_font: str, size: float, bold: bool = False, color: str = "000000") -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        el = OxmlElement("w:updateFields")
        el.set(qn("w:val"), "true")
        settings.append(el)


def add_field(paragraph, field_code: str, placeholder: str = "") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    for el in (begin, instr, separate):
        run = OxmlElement("w:r")
        run.append(el)
        paragraph._p.append(run)

    if placeholder:
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = placeholder
        run.append(text)
        paragraph._p.append(run)

    run = OxmlElement("w:r")
    run.append(end)
    paragraph._p.append(run)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.4)

    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(12)

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = "SimHei"
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
    heading1.font.size = Pt(16)
    heading1.font.bold = True

    heading2 = doc.styles["Heading 2"]
    heading2.font.name = "SimHei"
    heading2._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
    heading2.font.size = Pt(14)
    heading2.font.bold = True

    heading3 = doc.styles["Heading 3"]
    heading3.font.name = "SimHei"
    heading3._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
    heading3.font.size = Pt(12)
    heading3.font.bold = True

    for style_name, size in (("TOC 1", 12), ("TOC 2", 11.5), ("TOC 3", 11)):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = "SimSun"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        style.font.size = Pt(size)


def add_cover(doc: Document, title: str, subtitle: str, doc_type: str, doc_date: str | None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(126)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(title)
    set_run_font(run, "SimHei", "SimHei", 24, bold=True)

    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(18)
        run2 = p2.add_run(subtitle)
        set_run_font(run2, "SimHei", "SimHei", 18, bold=True)

    if doc_type:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(10)
        run3 = p3.add_run(doc_type)
        set_run_font(run3, "SimSun", "SimSun", 12, color="666666")

    if doc_date:
        p4 = doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.paragraph_format.space_before = Pt(168)
        run4 = p4.add_run(f"编制日期：{doc_date}")
        set_run_font(run4, "SimSun", "SimSun", 12)

    doc.add_page_break()


def add_toc(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("目录")
    set_run_font(run, "SimHei", "SimHei", 18, bold=True)

    toc = doc.add_paragraph()
    toc.paragraph_format.space_after = Pt(0)
    toc.paragraph_format.line_spacing = 1.2
    add_field(toc, 'TOC \\o "1-2" \\h \\z \\u', "目录将在更新域后生成")

    doc.add_page_break()


def add_inline_runs(paragraph, text: str) -> None:
    text = clean_inline(text)
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, "SimSun", "SimSun", 12, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, "SimSun", "SimSun", 12)


def parse_table_rows(table_lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for raw in table_lines:
        if re.match(r"^\|(?:\s*:?-+:?\s*\|)+\s*$", raw):
            continue
        rows.append([clean_inline(cell) for cell in raw.strip().strip("|").split("|")])
    return rows


def add_heading(doc: Document, level: int, text: str) -> None:
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    size_map = {1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph(style=style_map[level])
    p.paragraph_format.space_before = Pt(10 if level == 3 else 14)
    p.paragraph_format.space_after = Pt(5 if level == 3 else 6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, "SimHei", "SimHei", size_map[level], bold=True)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.84)
    p.paragraph_format.line_spacing = 1.72
    p.paragraph_format.space_after = Pt(5)
    add_inline_runs(p, text)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.25
            value = row[c_idx] if c_idx < len(row) else ""
            run = p.add_run(value)
            if r_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run_font(run, "SimHei", "SimHei", 11, bold=True)
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "D9E2F3")
                tc_pr.append(shd)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_run_font(run, "SimSun", "SimSun", 11)
    doc.add_paragraph()


def add_page_numbers(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_field(footer, "PAGE", "1")


def build(input_path: Path, output_path: Path, title: str | None, subtitle: str, doc_type: str, doc_date: str | None) -> None:
    lines = input_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    add_update_fields_on_open(doc)

    derived_title = title
    if not derived_title:
        derived_title = next((clean_inline(line[2:]) for line in lines if line.startswith("# ")), input_path.stem)

    add_cover(doc, derived_title, subtitle, doc_type, doc_date)
    add_toc(doc)

    table_lines: list[str] = []
    consumed_title = False

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("# ") and not consumed_title:
            consumed_title = True
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines.append(stripped)
            continue

        if table_lines:
            add_table(doc, parse_table_rows(table_lines))
            table_lines = []

        if not stripped or stripped == "---":
            continue

        if stripped.startswith("## "):
            add_heading(doc, 1, clean_inline(stripped[3:]))
        elif stripped.startswith("### "):
            add_heading(doc, 2, clean_inline(stripped[4:]))
        elif stripped.startswith("#### "):
            add_heading(doc, 3, clean_inline(stripped[5:]))
        else:
            add_body(doc, stripped)

    if table_lines:
        add_table(doc, parse_table_rows(table_lines))

    add_page_numbers(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build a formal Chinese DOCX from a Markdown manual.")
    parser.add_argument("--input", required=True, help="Input Markdown file path")
    parser.add_argument("--output", required=True, help="Output DOCX file path")
    parser.add_argument("--title", help="Document title for cover")
    parser.add_argument("--subtitle", default="软件说明书", help="Subtitle for cover")
    parser.add_argument("--doc-type", default="网页端功能说明文档", help="Small cover line")
    parser.add_argument("--doc-date", help="Document date shown on cover")
    args = parser.parse_args()

    build(
        input_path=Path(args.input),
        output_path=Path(args.output),
        title=args.title,
        subtitle=args.subtitle,
        doc_type=args.doc_type,
        doc_date=args.doc_date,
    )


if __name__ == "__main__":
    main()
